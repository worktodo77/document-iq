"""Generate the synthetic fixture corpus. No client data, ever.

Run it directly or import :func:`build`. The output is deterministic given the
same libraries: every timestamp, author and ID that a container format would
otherwise fill in from the clock is pinned to :data:`FIXED_TIMESTAMP`.

The generated files are NOT committed. PDF and OOXML containers embed producer
strings and creation times that vary by library version, so committing them
would put a binary blob in the repo that a dependency bump silently invalidates
— and the whole point of the corpus is that it is reproducible from this
script. ``tests/fixtures/generated/`` is gitignored; the tests build it on
demand.

Coverage is deliberate, and the mixed native+scanned PDF is the one the §11
reuse audit flagged as untested in the original extractor.
"""

from __future__ import annotations

import csv
import hashlib
import datetime
import io
import shutil
import time
import zipfile
from pathlib import Path

FIXED_TIMESTAMP = datetime.datetime(2024, 7, 16, 9, 30, 0)
"""Pinned so a rebuild of the fixtures does not change their bytes."""

_ZIP_DATE = (2024, 7, 16, 9, 30, 0)

HERE = Path(__file__).resolve().parent
OUT = HERE / "generated"


# ---------------------------------------------------------------------------
# PDFs
# ---------------------------------------------------------------------------


def _pin_ooxml(path: Path) -> None:
    """Rewrite an OOXML file with fixed zip member timestamps.

    DOCX and XLSX are zip containers, and python-docx / openpyxl stamp each
    member with the current time. Pinning ``core_properties`` fixes the
    metadata *inside* the parts and leaves the container varying, so two builds
    of the same fixture still differ — and because the file's SHA-256 is a Doc
    ID input, that propagates into the index, the ledger and every downstream
    artifact. The corpus then looks nondeterministic when only its inputs were.

    Member ORDER is preserved. OOXML readers rely on it — rewriting these
    archives in sorted order produces a file python-docx refuses to open — so
    determinism here comes from pinning the timestamps, never from reordering.
    """
    import xml.etree.ElementTree as ET

    stamp = FIXED_TIMESTAMP.strftime("%Y-%m-%dT%H:%M:%SZ")
    _DC = "{http://purl.org/dc/terms/}"

    def pin_core(data: bytes) -> bytes:
        # Parsed, not regexed. The obvious regex over `<dcterms:created …>`
        # also matches the `dcterms:W3CDTF` inside the element's own xsi:type
        # attribute, which silently corrupts the part.
        root = ET.fromstring(data)
        for tag in ("created", "modified"):
            for el in root.iter(_DC + tag):
                el.text = stamp
        for prefix, uri in (
            ("cp", "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"),
            ("dc", "http://purl.org/dc/elements/1.1/"),
            ("dcterms", "http://purl.org/dc/terms/"),
            ("dcmitype", "http://purl.org/dc/dcmitype/"),
            ("xsi", "http://www.w3.org/2001/XMLSchema-instance"),
        ):
            ET.register_namespace(prefix, uri)
        return ET.tostring(root, encoding="UTF-8", xml_declaration=True)

    with zipfile.ZipFile(path) as zin:
        payload = [(i.filename, zin.read(i.filename)) for i in zin.infolist()]

    # openpyxl rewrites dcterms:modified with the clock at save time whatever
    # `wb.properties.modified` was set to, so the value has to be pinned after
    # the fact rather than before it.
    payload = [
        (name, pin_core(data) if name == "docProps/core.xml" else data)
        for name, data in payload
    ]
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for name, data in payload:
            info = zipfile.ZipInfo(name, date_time=_ZIP_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            info.external_attr = 0o600 << 16
            zout.writestr(info, data)


def _pdf_canvas(path: Path):
    from reportlab.pdfgen import canvas

    # invariant=1 is the load-bearing argument: without it reportlab stamps a
    # CreationDate and a random document ID into every PDF, so two builds of
    # the fixture corpus differ and the self-test's reported corpus hash means
    # nothing across sessions. Setting the title/author/producer below does NOT
    # achieve this, which is what an earlier comment here claimed.
    #
    # The blast radius was wider than the PDFs: 14_transmittal.eml embeds one,
    # and 10_misnamed.docx IS one under a wrong extension, so a single
    # unpinned timestamp moved six of fourteen fixture files.
    c = canvas.Canvas(str(path), invariant=1)
    c.setTitle("DocIQ synthetic fixture")
    c.setAuthor("DocIQ fixture generator")
    c.setSubject("synthetic")
    c.setProducer("DocIQ fixtures")
    return c


def _text_page(c, lines: list[str]) -> None:
    y = 760
    for line in lines:
        c.drawString(60, y, line)
        y -= 18
    c.showPage()


def _image_page(c, lines: list[str]) -> None:
    """A page whose text is DRAWN, not typed — no text layer, so it can only be
    read by OCR. Large, high-contrast, plain glyphs: the fixture must exercise
    the OCR route, not benchmark the engine."""
    from reportlab.lib.utils import ImageReader
    from PIL import Image, ImageDraw

    img = Image.new("L", (1240, 1754), 255)
    d = ImageDraw.Draw(img)
    y = 120
    for line in lines:
        # Default bitmap font scaled up: legible to OCR without shipping a TTF.
        tile = Image.new("L", (620, 40), 255)
        ImageDraw.Draw(tile).text((4, 8), line, fill=0)
        img.paste(tile.resize((1116, 72), Image.LANCZOS), (60, y))
        y += 120
    c.drawImage(ImageReader(img), 0, 0, width=595, height=842)
    c.showPage()


def native_pdf(path: Path) -> None:
    c = _pdf_canvas(path)
    _text_page(c, ["MONTHLY PROGRESS REPORT", "Period: 2024-07-01 to 2024-07-31",
                   "Prepared by: Synthetic Contractor Ltd"])
    _text_page(c, ["2. PROGRESS NARRATIVE",
                   "Piling completed 16 July 2024.",
                   "Steel erection commenced 22/07/2024."])
    c.save()


def scanned_pdf(path: Path) -> None:
    c = _pdf_canvas(path)
    _image_page(c, ["SITE INSTRUCTION 014", "DATED 2024-07-16"])
    _image_page(c, ["ISSUED TO CONTRACTOR"])
    c.save()


def mixed_pdf(path: Path) -> None:
    """Native page, scanned page, native page — the audit's untested case."""
    c = _pdf_canvas(path)
    _text_page(c, ["TRANSMITTAL 2024-07-16",
                   "Attached: one scanned instruction sheet."])
    _image_page(c, ["SITE INSTRUCTION 015", "DATED 2024-07-17"])
    _text_page(c, ["End of transmittal.",
                   "Acknowledged 18 July 2024."])
    c.save()


def empty_page_pdf(path: Path) -> None:
    """Page 2 is genuinely blank: no text layer and no image, so neither route
    recovers anything. It must still be page 2 of 3."""
    c = _pdf_canvas(path)
    _text_page(c, ["COVER SHEET — CONTRACT ADMINISTRATION FILE",
                   "Document dated 2024-07-16, issued to the Engineer."])
    c.showPage()  # a page with nothing on it at all
    _text_page(c, ["APPENDIX A — SCHEDULE OF ATTACHMENTS",
                   "Nothing further was appended to this transmittal."])
    c.save()


# ---------------------------------------------------------------------------
# Office / text formats
# ---------------------------------------------------------------------------


def docx(path: Path) -> None:
    import docx as _docx

    d = _docx.Document()
    d.add_paragraph("CORRESPONDENCE")
    d.add_paragraph("Dear Sir, ref. our letter of 16 July 2024.")
    t = d.add_table(rows=2, cols=2)
    t.cell(0, 0).text = "Item"
    t.cell(0, 1).text = "Date"
    t.cell(1, 0).text = "Notice of delay"
    t.cell(1, 1).text = "2024-07-16"
    d.core_properties.created = FIXED_TIMESTAMP
    d.core_properties.modified = FIXED_TIMESTAMP
    d.core_properties.author = "DocIQ fixtures"
    d.core_properties.last_modified_by = "DocIQ fixtures"
    d.core_properties.revision = 1
    d.save(str(path))
    _pin_ooxml(path)


def xlsx(path: Path) -> None:
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Register"
    ws.append(["Ref", "Description", "Date"])
    ws.append(["RFI-001", "Foundation query", datetime.date(2024, 7, 16)])
    ws2 = wb.create_sheet("Empty")
    ws2["A1"] = None
    wb.properties.created = FIXED_TIMESTAMP
    wb.properties.modified = FIXED_TIMESTAMP
    wb.properties.creator = "DocIQ fixtures"
    wb.properties.lastModifiedBy = "DocIQ fixtures"
    wb.save(str(path))
    _pin_ooxml(path)


def csv_file(path: Path) -> None:
    buf = io.StringIO(newline="")
    w = csv.writer(buf, lineterminator="\n")
    w.writerow(["Ref", "Description", "Date"])
    w.writerow(["NCR-007", "Weld porosity; rework 2024-07-18", "2024-07-16"])
    path.write_text(buf.getvalue(), encoding="utf-8", newline="\n")


def txt_file(path: Path) -> None:
    """Carries the adversarial normalization input on purpose: mixed CRLF/CR,
    a run of NBSPs, stacked zero-width characters and a four-blank-line gap.
    The fixture corpus is where a normalization regression should surface."""
    payload = (
        "DAILY LOG\r\n"
        "Date: 2024-07-16\r"
        "Crew:   12 operatives   \n"
        "​​﻿Delay‍ noted.\n"
        "\n\n\n\n"
        "End of log.\t\n"
    )
    path.write_text(payload, encoding="utf-8", newline="")


def eml_file(path: Path) -> None:
    """RFC-822 message. Written by hand rather than through ``email`` so the
    bytes are pinned — a library's own Message-ID and boundary generation is
    clock- and random-seeded."""
    lines = [
        "From: engineer@example.com",
        "To: contractor@example.com",
        "Cc: pm@example.com",
        "Subject: Notice of Delay - Area 200",
        "Date: Tue, 16 Jul 2024 09:30:00 +0000",
        "Message-ID: <fixture-0001@example.com>",
        "MIME-Version: 1.0",
        "Content-Type: text/plain; charset=utf-8",
        "",
        "Please treat this as notice under clause 20.1.",
        "The delay commenced 2024-07-16 and is ongoing.",
        "",
    ]
    path.write_bytes("\r\n".join(lines).encode("utf-8"))


def eml_with_attachment(path: Path, attachment: Path) -> None:
    """A message carrying a PDF — §3's "attachments extracted as child documents".

    The corpus needs this case in its own right, not only in a unit test. Email
    attachment expansion has a *different producer* in the walker from archive
    expansion, so a ZIP proving the container path proves nothing about this one;
    and until the attachment case is in the corpus, the determinism proof and the
    self-test never touch it. "The corpus doesn't exercise it" is not a reason to
    leave it out — it is a reason to put it in.

    Written by hand for the same reason as :func:`eml_file`: a library's own
    boundary and Message-ID generation is clock- and random-seeded, and these
    bytes have to be identical on every regeneration.
    """
    import base64

    payload = base64.b64encode(attachment.read_bytes()).decode("ascii")
    body = "\r\n".join(payload[i:i + 76] for i in range(0, len(payload), 76))
    lines = [
        "From: engineer@example.com",
        "To: contractor@example.com",
        "Subject: Transmittal 2024-07-18 - one attachment",
        "Date: Thu, 18 Jul 2024 11:00:00 +0000",
        "Message-ID: <fixture-0002@example.com>",
        "MIME-Version: 1.0",
        'Content-Type: multipart/mixed; boundary="dociq-fixture-boundary"',
        "",
        "--dociq-fixture-boundary",
        "Content-Type: text/plain; charset=utf-8",
        "",
        "Please find the monthly report attached.",
        "Issued 2024-07-18 under clause 20.1.",
        "",
        "--dociq-fixture-boundary",
        'Content-Type: application/pdf; name="attached_report.pdf"',
        "Content-Transfer-Encoding: base64",
        'Content-Disposition: attachment; filename="attached_report.pdf"',
        "",
        body,
        "",
        "--dociq-fixture-boundary--",
        "",
    ]
    path.write_bytes("\r\n".join(lines).encode("utf-8"))


def nested_zip(path: Path, inner_sources: list[Path]) -> None:
    """A ZIP holding a ZIP — exercises the depth guard and child ordering."""
    inner = io.BytesIO()
    with zipfile.ZipFile(inner, "w", zipfile.ZIP_DEFLATED) as zf:
        for src in inner_sources:
            info = zipfile.ZipInfo(src.name, date_time=_ZIP_DATE)
            info.compress_type = zipfile.ZIP_DEFLATED
            zf.writestr(info, src.read_bytes())
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        info = zipfile.ZipInfo("readme.txt", date_time=_ZIP_DATE)
        zf.writestr(info, b"Archive produced 2024-07-16.\n")
        info = zipfile.ZipInfo("inner.zip", date_time=_ZIP_DATE)
        zf.writestr(info, inner.getvalue())


def tier2_file(path: Path) -> None:
    """A legacy .doc — D-02's Tier-2 case. The bytes are an OLE header so a
    content sniff would recognize the container; the point is that DocIQ never
    tries, because §3 lists it rather than extracting it."""
    path.write_bytes(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1" + b"\x00" * 504)


def misnamed_pdf(path: Path) -> None:
    """PDF bytes under a .docx name — the content-sniffing recovery case that
    litigation productions actually deliver."""
    tmp = path.with_suffix(".sniff.tmp")
    native_pdf(tmp)
    path.write_bytes(tmp.read_bytes())
    tmp.unlink()


# ---------------------------------------------------------------------------


def _generator_stamp() -> str:
    """Identity of THIS generator, readable from a checkout or from a bundle.

    ``sha256(Path(__file__).read_bytes())`` is the natural expression and it is
    wrong inside a PyInstaller build: ``__file__`` names a path in the archive
    that does not exist on disk, so the packaged self-test died with
    ``FileNotFoundError`` before it built a single fixture. This is the fourth
    member of one class — every place in the tree that derives a *path* from
    ``__file__`` and expects a real file there. The other three were
    ``branding.palette``'s brand directory, ``ingest.extract``'s OCR model
    directory and ``verify.determinism``'s subprocess ``PYTHONPATH``; this one
    was missed on the first sweep because the sweep read ``src/`` and this file
    lives under ``tests/``. Recorded rather than quietly fixed, because the
    lesson is about the sweep, not about the line.

    Frozen, the module's own compiled code stands in for its source bytes. It
    changes when the generator changes, which is the whole property the stamp
    needs; it is simply not the same VALUE as the source hash, so a corpus
    built by the frozen build and one built from a checkout do not share a
    completion marker. They also never share a directory, so nothing rebuilds
    that would not have rebuilt anyway.
    """
    try:
        return hashlib.sha256(Path(__file__).read_bytes()).hexdigest()[:16]
    except OSError:
        import importlib.util
        import marshal

        spec = importlib.util.find_spec(__name__)
        code = spec.loader.get_code(__name__)  # type: ignore[union-attr]
        return hashlib.sha256(marshal.dumps(code)).hexdigest()[:16]


def build(out: Path | None = None) -> Path:
    """(Re)generate the corpus and return its root.

    Concurrency-safe, and it has to be. Two pytest sessions in one worktree
    both import ``conftest``, and an unguarded rebuild rewrites files a
    concurrent session is mid-read of — which changes extracted bytes and so
    changes the content hash, failing the determinism tests intermittently
    while the product is correct. That reads as a determinism defect, which is
    the most expensive kind of false alarm this project can raise.

    The guard is a directory-creation lock (atomic on Windows and POSIX) plus a
    completion marker, so the second session waits and then reuses the corpus
    rather than rebuilding it.
    """
    out = Path(out) if out is not None else OUT
    src = out / "matter"
    lock = out / ".build.lock"
    done = out / ".build.complete"
    # The marker records WHICH generator built the corpus, not merely that one
    # did. Keyed on "ok" alone, editing this file left a stale corpus in place
    # forever and the tests silently kept asserting against the old bytes --
    # which is exactly how a determinism fix here first looked like a product
    # regression.
    stamp = _generator_stamp()
    out.mkdir(parents=True, exist_ok=True)

    for _ in range(600):  # ~60s; the build itself takes a few seconds
        if done.exists() and done.read_text(encoding='utf-8').strip() == stamp:
            return src
        try:
            lock.mkdir()
            break
        except FileExistsError:
            time.sleep(0.1)
    else:
        # Never block a test run on a lock left behind by a killed session.
        shutil.rmtree(lock, ignore_errors=True)
        lock.mkdir(exist_ok=True)

    try:
        _build_corpus(src)
        done.write_text(stamp, encoding="utf-8")
    finally:
        shutil.rmtree(lock, ignore_errors=True)
    return src


def _build_corpus(src: Path) -> Path:
    sub = src / "attachments"
    sub.mkdir(parents=True, exist_ok=True)

    native_pdf(src / "01_native_report.pdf")
    scanned_pdf(src / "02_scanned_instruction.pdf")
    mixed_pdf(src / "03_mixed_transmittal.pdf")
    empty_page_pdf(src / "04_empty_page.pdf")
    docx(src / "05_letter.docx")
    xlsx(src / "06_register.xlsx")
    csv_file(src / "07_ncr_log.csv")
    txt_file(src / "08_daily_log.txt")
    eml_file(src / "09_notice.eml")
    eml_with_attachment(src / "14_transmittal.eml", src / "01_native_report.pdf")
    tier2_file(src / "13_legacy.doc")
    misnamed_pdf(sub / "10_misnamed.docx")
    nested_zip(src / "11_production.zip",
               [src / "07_ncr_log.csv", src / "08_daily_log.txt",
                src / "13_legacy.doc"])
    # Same bytes as 07 under a second path — the duplicate-by-hash case.
    (sub / "12_ncr_log_copy.csv").write_bytes((src / "07_ncr_log.csv").read_bytes())
    return src


if __name__ == "__main__":  # pragma: no cover — developer entry point
    print(build())
